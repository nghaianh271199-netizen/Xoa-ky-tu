import streamlit as st
import re
import json
from docx import Document

st.set_page_config(page_title="Văn bản chuẩn hóa (OCR fix)", layout="wide")

# --- helper: extract full text from docx (paragraphs + tables + headers/footers)
def extract_text_from_docx(file_stream):
    doc = Document(file_stream)
    parts = []

    def para_text(p):
        # join runs to keep raw characters
        return "".join([r.text for r in p.runs]).strip()

    for p in doc.paragraphs:
        t = para_text(p)
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = para_text(p)
                    if t:
                        parts.append(t)

    for section in doc.sections:
        try:
            for p in section.header.paragraphs:
                t = para_text(p)
                if t:
                    parts.append(t)
            for p in section.footer.paragraphs:
                t = para_text(p)
                if t:
                    parts.append(t)
        except Exception:
            pass

    return "\n".join(parts)


# --- Detect if a string contains Vietnamese vowel characters (unicode-aware) ---
VOWEL_PATTERN = r"[aeiouyàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ]"
_vowel_re = re.compile(VOWEL_PATTERN, flags=re.I | re.UNICODE)


# === improved OCR merging:
# Try to greedily merge small tokens into plausible syllables:
def merge_ocr_tokens(text: str) -> str:
    if not text:
        return text

    # tokenise but keep punctuation as separate tokens
    tokens = re.findall(r"\w+|[^\w\s]", text, flags=re.UNICODE)
    out_tokens = []
    i = 0
    n = len(tokens)

    while i < n:
        tok = tokens[i]

        # if punctuation, just append
        if re.fullmatch(r'[^\w\s]', tok, flags=re.UNICODE):
            out_tokens.append(tok)
            i += 1
            continue

        # if token already contains a vowel, treat alone (likely a valid syllable/word)
        if _vowel_re.search(tok):
            out_tokens.append(tok)
            i += 1
            continue

        # otherwise try to merge with next 1 or 2 tokens to create candidate containing vowel
        merged = None
        # try length 3 (tok + tok1 + tok2)
        if i + 2 < n:
            cand = tokens[i] + tokens[i+1] + tokens[i+2]
            if not any(re.fullmatch(r'[^\w\s]', x, flags=re.UNICODE) for x in (tokens[i+1], tokens[i+2])) and _vowel_re.search(cand):
                merged = cand
                i += 3
        # try length 2 (tok + tok1)
        if merged is None and i + 1 < n:
            cand = tokens[i] + tokens[i+1]
            if not re.fullmatch(r'[^\w\s]', tokens[i+1], flags=re.UNICODE) and _vowel_re.search(cand):
                merged = cand
                i += 2

        if merged:
            out_tokens.append(merged)
        else:
            # nothing to merge -> append tok
            out_tokens.append(tok)
            i += 1

    # Reconstruct text: put a space between tokens except before punctuation and after opening punctuation
    res_parts = []
    for idx, t in enumerate(out_tokens):
        res_parts.append(t)
        if idx + 1 < len(out_tokens):
            nxt = out_tokens[idx + 1]
            if not re.fullmatch(r'[^\w\s]', nxt, flags=re.UNICODE):  # next is not punctuation
                res_parts.append(" ")

    return "".join(res_parts)


# === Fix missing spacing when two words are glued: "ởđó" -> "ở đó"
def fix_missing_spacing(text: str) -> str:
    # add space between vowel-with-diacritic and ascii letter, or ascii letter and vowel-with-diacritic
    vowels = "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ"
    pattern1 = rf'([{vowels}])([A-Za-z])'
    pattern2 = rf'([A-Za-z])([{vowels}])'
    text = re.sub(pattern1, r'\1 \2', text, flags=re.UNICODE)
    text = re.sub(pattern2, r'\1 \2', text, flags=re.UNICODE)
    return text


# === Remove redundant spaces and collapse repeated dots
def fix_basic_spacing(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\.{2,}', '.', text)
    return text.strip()


# === Full normalize pipeline (applies all fixes) ===
def normalize_text(text: str) -> str:
    if not text:
        return ""

    # 1) remove undesirable characters but keep basic punctuation . , ; : ? ! ( ) and newlines
    text = re.sub(r'[“”\"\'\*\~\^\%\$\#\@\[\]\{\}\<\>\\\/\=\+]', '', text)

    # 2) normalize hyphen types -> dot (per requirement)
    text = text.replace('-', '.')

    # 3) collapse multiple dots
    text = re.sub(r'\.{2,}', '.', text)

    # 4) lower-case the whole text (spec: bold -> lowercase; we keep overall lowercase to be consistent)
    text = text.lower()

    # 5) basic spacing cleanup
    text = fix_basic_spacing(text)

    # 6) merge OCR-separated letters/syllables
    text = merge_ocr_tokens(text)

    # 7) after merging, fix leftover small spacing artifacts
    text = fix_basic_spacing(text)

    # 8) fix glued words (missing spacing)
    text = fix_missing_spacing(text)

    # 9) final cleanup
    text = fix_basic_spacing(text)

    return text


# === copy button using JS (works in Streamlit) ===
def js_copy_button(text: str, label: str = "📋 Copy toàn bộ"):
    # safe JSON encode the text to embed in JS
    encoded = json.dumps(text)
    html = f"""
        <button id="copy-btn">{label}</button>
        <script>
        const btn = document.getElementById("copy-btn");
        btn.addEventListener("click", async () => {{
            try {{
                await navigator.clipboard.writeText({encoded});
                btn.innerText = "✅ Đã copy";
            }} catch (e) {{
                alert("Không thể copy tự động — vui lòng bôi đen và Ctrl+C");
            }}
        }});
        </script>
    """
    st.markdown(html, unsafe_allow_html=True)


# ------------------ Streamlit UI ------------------
st.title("📄 Công cụ Chuẩn hóa Văn bản (OCR-fix nâng cao)")
st.write(
    "Upload DOCX hoặc dán văn bản rồi bấm **⚙️ Xử lý văn bản**. "
    "Quy tắc: loại bỏ ngoặc kép và ký tự đặc biệt, '-' -> '.', gom nhiều '.' thành 1, "
    "sửa lỗi OCR (ký tự/syllable bị tách), chèn khoảng trắng nếu từ bị dính."
)

mode = st.radio("Phương thức nhập:", ["📂 Tải file DOCX", "⌨️ Nhập văn bản"])
input_text = ""

if mode.startswith("📂"):
    uploaded = st.file_uploader("Chọn file .docx", type=["docx"])
    if uploaded:
        try:
            input_text = extract_text_from_docx(uploaded)
        except Exception as e:
            st.error("Lỗi khi đọc file .docx: " + str(e))

else:
    input_text = st.text_area("Dán hoặc nhập văn bản ở đây:", height=300)

# xử lý khi bấm nút
if st.button("⚙️ Xử lý văn bản"):
    if not input_text or not input_text.strip():
        st.warning("Vui lòng tải lên hoặc nhập văn bản trước khi xử lý.")
    else:
        st.subheader("📌 Văn bản gốc")
        st.text_area("Gốc", value=input_text, height=200)

        processed = normalize_text(input_text)

        st.subheader("✅ Văn bản đã chuẩn hóa")
        st.text_area("Kết quả", value=processed, height=300)

        # JS copy button
        js_copy_button(processed)

        st.info("Ghi chú: thuật toán là heuristic — nếu còn chỗ chưa đúng, mình có thể bổ sung quy tắc hoặc dùng bộ tách từ tiếng Việt (underthesea/pyvi) để cải thiện chính xác hơn.")
